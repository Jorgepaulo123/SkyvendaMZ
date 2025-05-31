from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_paragraph_centered(cell, text, bold=False):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold

def create_invoice():
    doc = Document()

    # Cabeçalho
    doc.add_paragraph("Aliyah Interprise").runs[0].bold = True
    doc.add_paragraph("NUIT: 401541319")
    doc.add_paragraph("Av. Júlio Nyerere, Bairro Central Mahory - Mueda, Cabo Delgado")
    doc.add_paragraph("Cel: +258 842021739 / 825070888")
    doc.add_paragraph("NIASSA")

    doc.add_paragraph("FACTURA Nº: 00227")

    doc.add_paragraph("Cliente: ________________________________")
    doc.add_paragraph("Morada: ________________________________")
    doc.add_paragraph("NUIT: _____________    Data: ___/___/20___")

    # Tabela principal
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['QTD', 'DESIGNAÇÃO', 'PREÇO UNITÁRIO', 'TOTAL']
    for i, text in enumerate(headers):
        add_paragraph_centered(hdr_cells[i], text, bold=True)

    # 15 linhas em branco
    for _ in range(15):
        row_cells = table.add_row().cells
        for cell in row_cells:
            cell.text = ""

    doc.add_paragraph()

    # Motivo do IVA
    table_iva = doc.add_table(rows=1, cols=1)
    table_iva.style = 'Table Grid'
    table_iva.cell(0, 0).text = "Motivo justificativo de não pagamento do IVA: _________________________________"

    # Resumo de valores
    table_valores = doc.add_table(rows=3, cols=2)
    table_valores.style = 'Table Grid'
    table_valores.cell(0, 0).text = "SUB-TOTAL:"
    table_valores.cell(1, 0).text = "IVA 16%:"
    table_valores.cell(2, 0).text = "TOTAL:"
    for i in range(3):
        table_valores.cell(i, 1).text = "...................."

    # Forma de pagamento
    doc.add_paragraph("Pagamento: [ ] Dinheiro  [ ] Cheque Nº ______  [ ] Outros ______________")

    # Dados bancários
    doc.add_paragraph("NUMERO DE CONTA BCI: 24106672910001 - NIB: 0008.0000.4106472910113")
    doc.add_paragraph("IBAN: MZ59 0008.0000.4106472910113")

    doc.add_paragraph("Atenciosamente, ________________________")

    # Salvar o ficheiro
    doc.save("factura_modelo.docx")
    print("Factura criada como 'factura_modelo.docx'")

create_invoice()
